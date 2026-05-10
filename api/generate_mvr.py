"""
generate_mvr_v3.py
Fills the approved Shaigan MVR template with data from Excel.
Usage: python generate_mvr_v3.py <excel.xlsx> [output.docx]
Requires approved_file.docx in the same folder.
"""
from docx.shared import Inches
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx import Document
import numpy as np
import matplotlib.pyplot as plt
import sys
import shutil
import os
import copy
import re
import tempfile
import zipfile
import openpyxl
import matplotlib
matplotlib.use('Agg')

# ── helpers ────────────────────────────────────────────────────────────────


def fmt(v, dec=2):
    if v is None or v == "":
        return ""
    try:
        return f"{float(v):.{dec}f}"
    except:
        return str(v)


def fmt0(v):
    if v is None or v == "":
        return ""
    try:
        return str(int(round(float(v))))
    except:
        return str(v)


def set_run_text(run, text):
    """Replace text in a run preserving all formatting."""
    run.text = str(text) if text is not None else ""


def replace_in_para(para, old, new):
    """Replace all occurrences of old with new across runs in a paragraph."""
    full = "".join(r.text for r in para.runs)
    if old not in full:
        return False
    new_full = full.replace(old, str(new))
    if para.runs:
        para.runs[0].text = new_full
        for r in para.runs[1:]:
            r.text = ""
    return True


def replace_bold_in_cell(cell, old_text, new_text):
    """Replace text in cell making the new_text bold, keeping surrounding text normal."""
    from docx.oxml.ns import qn as _qn
    from lxml import etree
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    for para in cell.paragraphs:
        full = "".join(r.text for r in para.runs)
        if old_text not in full:
            continue
        parts = full.split(old_text)
        p_el = para._p
        for r in para.runs:
            p_el.remove(r._r)

        def make_run(text, bold=False):
            r_el = etree.SubElement(p_el, f'{{{W}}}r')
            rpr = etree.SubElement(r_el, f'{{{W}}}rPr')
            sz = etree.SubElement(rpr, f'{{{W}}}sz')
            sz.set(f'{{{W}}}val', '22')
            szcs = etree.SubElement(rpr, f'{{{W}}}szCs')
            szcs.set(f'{{{W}}}val', '22')
            if bold:
                etree.SubElement(rpr, f'{{{W}}}b')
                etree.SubElement(rpr, f'{{{W}}}bCs')
            t_el = etree.SubElement(r_el, f'{{{W}}}t')
            t_el.text = text
            t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            return r_el
        for i, part in enumerate(parts):
            if part:
                make_run(part, bold=False)
            if i < len(parts) - 1:
                make_run(new_text, bold=True)
        break


def replace_cell_text(cell, new_text, preserve_bold=None):
    """Replace all text in a cell's first paragraph runs."""
    for para in cell.paragraphs:
        if para.runs:
            first = para.runs[0]
            first.text = str(new_text) if new_text is not None else ""
            if preserve_bold is not None:
                first.bold = preserve_bold
            for r in para.runs[1:]:
                r.text = ""
            break


def replace_tc(row, tc_idx, new_text, bold=None):
    """Replace text in actual tc element by index using Run API (survives doc.save())."""
    from docx.oxml.ns import qn as _qn
    from docx.text.run import Run
    tcs = row._tr.findall(_qn('w:tc'))
    if tc_idx >= len(tcs):
        return
    tc = tcs[tc_idx]
    paras = tc.findall(_qn('w:p'))
    if not paras:
        return
    runs = paras[0].findall(_qn('w:r'))
    if not runs:
        return
    run_obj = Run(runs[0], None)
    run_obj.text = str(new_text) if new_text is not None else ""
    if bold is not None:
        run_obj.bold = bold
    for r in runs[1:]:
        Run(r, None).text = ""


def cell_text(cell):
    return "".join(r.text for p in cell.paragraphs for r in p.runs)


def find_run_with(para, text):
    full = "".join(r.text for r in para.runs)
    return text in full

# ── Excel reader ───────────────────────────────────────────────────────────


def read_excel(path):
    wb = openpyxl.load_workbook(path, data_only=True)
    wi = wb["Product Info"]
    info = {}
    for row in wi.iter_rows(min_row=1, values_only=True):
        if row[0] and row[1] is not None:
            info[str(row[0]).strip()] = row[1]

    # Product info fields
    PN = str(info.get("Product Name", ""))
    SMP = str(info.get("Starting Material/Product", "sm")).strip().lower()
    AI = str(info.get("Active Ingredient", ""))
    MVL = str(info.get("MVL Number", "") or "")
    MVER = str(info.get("MVER Number", "") or "")
    TEST = f"Assay of {AI}"
    TP = str(info.get("Testing Procedure", ""))
    HID = str(info.get("HPLC ID", ""))
    WS = str(info.get("Working Standard", ""))
    BATCH = str(info.get("Batch #", ""))
    PUR = info.get("Purity", 99.66)

    # For Product: read batch 2 and 3 from rows 12 and 13 by position
    is_product = "p" in SMP or "product" in SMP
    if is_product:
        BATCH2 = str(wi.cell(row=12, column=2).value or "").strip()
        BATCH3 = str(wi.cell(row=13, column=2).value or "").strip()
        batches = [b for b in [BATCH, BATCH2, BATCH3] if b]
        BATCH_STR = ", ".join(batches)
    else:
        BATCH_STR = BATCH

    # Determine report number and name token
    RPT_NO = MVL if MVL else MVER
    NAME = AI if ("sm" in SMP or "starting" in SMP) else PN

    ws = wb["Validation Data"]
    def cv(c, r): return ws.cell(row=r, column=c).value

    def fv(c, r, dec=2):
        v = cv(c, r)
        if v is None or v == "":
            return ""
        try:
            return round(float(v), dec)
        except:
            return str(v)

    # Locate sections
    sec = {}
    for r in range(1, ws.max_row + 1):
        a = str(ws.cell(row=r, column=1).value or "").strip()
        if a.startswith("1. System"):
            sec["assay"] = r
        elif a.startswith("2. Lin"):
            sec["lin"] = r
        elif a.startswith("3. Acc"):
            sec["acc"] = r
        elif a.startswith("4. Int"):
            sec["ip"] = r
        elif a.startswith("5. For"):
            sec["deg"] = r
        elif a.startswith("6. Rob"):
            sec["rob"] = r
        elif a.startswith("7. LOD"):
            sec["lod"] = r
        elif a.startswith("8. Lin"):
            sec["chart"] = r

    # System Suitability
    d1_ar = sec["assay"] + 2
    d1 = {
        "areas": [cv(c, d1_ar) for c in range(2, 8)],
        "av":  fv(2, d1_ar+1, 2),
        "sd":  fv(2, d1_ar+2, 2),
        "rsd": fv(2, d1_ar+3, 2),
    }

    # Linearity
    ls = sec["lin"] + 2
    lin = []
    conc_map = {80: 0.122, 90: 0.135, 100: 0.150, 110: 0.165, 120: 0.182}
    for level, ac in [(80, 2), (90, 4), (100, 6)]:
        lin.append(
            {"pct": level, "conc": conc_map[level], "av": fv(ac, ls+3, 0)})
    ls2 = ls + 7
    for level, ac in [(110, 2), (120, 4)]:
        lin.append(
            {"pct": level, "conc": conc_map[level], "av": fv(ac, ls2+3, 0)})

    # Accuracy
    as_ = sec["acc"] + 2
    acc = []
    for level, ac, rc in [(80, 2, 3), (100, 4, 5), (120, 6, 7)]:
        acc.append({
            "pct": level, "conc": conc_map[level],
            "areas":  [cv(ac, as_+i) for i in range(3)],
            "recs":   [fv(rc, as_+i, 2) for i in range(3)],
            "rsd":    fv(ac, as_+4, 2),
            "av_rec": fv(rc, as_+3, 2),
        })

    # IP
    ip_start = sec["ip"] + 2
    ip = []
    for offset in [0, 7, 14]:
        r0 = ip_start + offset
        ip.append({
            "areas":   [cv(3, r0+i) for i in range(3)],
            "results": [fv(4, r0+i, 2) for i in range(3)],
            "avg_r":   fv(4, r0+3, 2),
            "rsd_r":   fv(4, r0+5, 2),
        })

    # Forced Degradation
    ds = sec["deg"] + 2
    deg = []
    for name in ["Acid Degradation", "Base Degradation", "H\u2082O\u2082 Degradation"]:
        deg.append({
            "name":    name,
            "areas":   [cv(3, ds+i) for i in range(2)],
            "results": [fv(4, ds+i, 2) for i in range(2)],
            "av_r":    fv(6, ds, 2),
        })
        ds += 3

    # Robustness
    rs = sec["rob"] + 2
    rob = []
    for i in range(0, 4, 2):
        rob.append({
            "flow":   str(cv(1, rs+i) or ""),
            "std_a":  cv(3, rs+i),
            "spl_a":  cv(3, rs+i+1),
            "result": fv(4, rs+i, 2),
        })

    # LOD/LOQ
    lr = sec["lod"] + 2
    lod = {
        "sd":    fv(5, lr, 2),
        "slope": fv(6, lr, 0),
        "ug":    fv(3, lr, 2),
        "area":  fv(4, lr, 0),
    }
    loq = {
        "ug":   fv(3, lr+1, 2),
        "area": fv(4, lr+1, 0),
    }

    chart_concs = [l["conc"] for l in lin]
    chart_areas = [l["av"] for l in lin]

    return dict(
        PN=PN, AI=AI, NAME=NAME, RPT_NO=RPT_NO, SMP=SMP,
        TEST=TEST, TP=TP, HID=HID, WS=WS,
        BATCH=BATCH, BATCH_STR=BATCH_STR, PUR=PUR,
        d1=d1, lin=lin, acc=acc, ip=ip, deg=deg, rob=rob,
        lod=lod, loq=loq,
        chart_concs=chart_concs, chart_areas=chart_areas,
    )


# ── Template filler ────────────────────────────────────────────────────────

def fill_template(data, template, output):
    shutil.copy(template, output)
    doc = Document(output)

    AI = data["AI"]
    PN = data["PN"]
    NAME = data["NAME"]
    SMP = data["SMP"]
    RPT = data["RPT_NO"]
    d1 = data["d1"]
    lin = data["lin"]
    acc = data["acc"]
    ip = data["ip"]
    deg = data["deg"]
    rob = data["rob"]
    lod = data["lod"]
    loq = data["loq"]
    is_product = "p" in SMP or "product" in SMP

    body = doc.element.body
    tbls = doc.tables
    paras = doc.paragraphs

    # ── P[001]: Report No ─────────────────────────────────────────────────
    for p in paras:
        if "Method Validation Report No:" in p.text:
            replace_in_para(p, "MVL-524", RPT)
            break

    # ── TBL[002]: Cover info table ────────────────────────────────────────
    tbl_cover = tbls[0]
    replace_cell_text(tbl_cover.cell(0, 1), AI)
    replace_cell_text(tbl_cover.cell(0, 3), data["TEST"])
    replace_cell_text(tbl_cover.cell(1, 1), data["TP"])

    # ── TBL[022]: Pre-verification ────────────────────────────────────────
    tbl_pre = tbls[2]
    cell_pre = tbl_pre.cell(0, 0)
    for para in cell_pre.paragraphs:
        full = "".join(r.text for r in para.runs)
        if "HPLC" in full and "QC-" in full:
            new_full = re.sub(r'HPLC\s+\S+', f'HPLC {data["HID"]}', full)
            if para.runs:
                para.runs[0].text = new_full
                for r in para.runs[1:]:
                    r.text = ""
        elif "Working Standard #" in full:
            replace_bold_in_cell(cell_pre, "Avanafil Working Standard # CPL/WS/25/AVN/038",
                                 f"{AI} Working Standard # {data['WS']}")
            # NO break — continue to reach batch line
        elif "Starting Material B #" in full:
            if is_product:
                replace_bold_in_cell(cell_pre, "Avanafil Starting Material B # 15048002-AVN",
                                     f"{PN} Batch # {data['BATCH_STR']}")
            else:
                replace_bold_in_cell(cell_pre, "Avanafil Starting Material B # 15048002-AVN",
                                     f"{AI} Starting Material B # {data['BATCH_STR']}")
            break

    # ── TBL[038]: Linearity data ──────────────────────────────────────────
    tbl_lin = tbls[7]
    for i, lv in enumerate(lin):
        row = tbl_lin.rows[i+1]
        replace_cell_text(row.cells[2], fmt0(lv["av"]))

    # ── Accuracy tables ───────────────────────────────────────────────────
    acc_tbls = [tbls[11], tbls[12], tbls[13]]
    for idx, lvl in enumerate(acc):
        t = acc_tbls[idx]
        for i in range(3):
            row = t.rows[i+1]
            replace_cell_text(row.cells[2], fmt0(lvl["areas"][i]))
            if i == 0:
                replace_cell_text(row.cells[3], fmt(lvl["rsd"], 2))
            replace_cell_text(row.cells[4], fmt(lvl["recs"][i], 2))

    # ── TBL[065]: Grand average ───────────────────────────────────────────
    tbl_gav = tbls[14]
    for i, lvl in enumerate(acc):
        row = tbl_gav.rows[i+1]
        replace_cell_text(row.cells[1], f"{lvl['av_rec']}%")

    # ── TBL[077]: Precision Repeatability ────────────────────────────────
    tbl_rep = tbls[17]
    for i, area in enumerate(d1["areas"]):
        replace_cell_text(tbl_rep.rows[i+1].cells[1], fmt0(area))
    replace_cell_text(tbl_rep.rows[7].cells[1], fmt0(d1["av"]))
    replace_cell_text(tbl_rep.rows[8].cells[1], fmt(d1["sd"], 2))
    replace_cell_text(tbl_rep.rows[9].cells[1], fmt(d1["rsd"], 2))

    # ── TBL[082]: Repeatability acceptance ───────────────────────────────
    tbl_rep_acc = tbls[18]
    replace_cell_text(tbl_rep_acc.rows[1].cells[1], fmt(d1["rsd"], 2))

    # ── TBL[021]: Within Days ─────────────────────────────────────────────
    tbl_wd = tbls[21]
    wd_data = [ip[0], ip[2]]
    day_lbls = ["Day 1", "Day 2"]
    row_idx = 2
    for di, ipd in enumerate(wd_data):
        for si in range(3):
            row = tbl_wd.rows[row_idx]
            replace_cell_text(row.cells[0], day_lbls[di])
            replace_cell_text(row.cells[1], fmt0(ipd["areas"][si]))
            replace_cell_text(row.cells[2], fmt(ipd["results"][si], 2))
            replace_cell_text(row.cells[3], fmt(ipd["avg_r"], 2))
            replace_cell_text(row.cells[4], fmt(ipd["rsd_r"], 2))
            row_idx += 1

    # ── TBL[022]: Within Days acceptance ─────────────────────────────────
    tbl_wd_acc = tbls[22]
    replace_cell_text(tbl_wd_acc.rows[1].cells[1], fmt(ip[0]["rsd_r"], 2))
    replace_cell_text(tbl_wd_acc.rows[2].cells[1], fmt(ip[2]["rsd_r"], 2))

    # ── TBL[104]: By Different Analyst ───────────────────────────────────
    tbl_an = tbls[24]
    an_data = [ip[0], ip[1]]
    row_idx = 1
    for ai, ipd in enumerate(an_data):
        for si in range(3):
            row = tbl_an.rows[row_idx]
            if si == 0:
                replace_cell_text(row.cells[0], str(ai+1))
                replace_cell_text(row.cells[3], fmt(ipd["avg_r"], 2))
                replace_cell_text(row.cells[4], fmt(ipd["rsd_r"], 2))
            replace_cell_text(row.cells[1], fmt0(ipd["areas"][si]))
            replace_cell_text(row.cells[2], fmt(ipd["results"][si], 2))
            row_idx += 1

    # ── TBL[107]: By Analyst acceptance ──────────────────────────────────
    tbl_an_acc = tbls[25]
    replace_cell_text(tbl_an_acc.rows[1].cells[1], fmt(ip[0]["rsd_r"], 2))
    replace_cell_text(tbl_an_acc.rows[2].cells[1], fmt(ip[1]["rsd_r"], 2))

    # ── TBL[027]: Robustness (merged cells) ──────────────────────────────
    tbl_rob = tbls[27]
    for ri, r in enumerate(rob):
        std_row = tbl_rob.rows[ri*2 + 3]
        spl_row = tbl_rob.rows[ri*2 + 4]
        replace_tc(std_row, 0, r["flow"])
        replace_tc(std_row, 1, fmt0(r["std_a"]))
        replace_tc(std_row, 2, fmt(r["result"], 2))
        replace_tc(spl_row, 0, "")
        replace_tc(spl_row, 1, fmt0(r["spl_a"]))
        replace_tc(spl_row, 2, "")

    # ── TBL[122]: Forced Degradation ─────────────────────────────────────
    tbl_deg = tbls[29]
    row_idx = 1
    for d in deg:
        for si in range(2):
            row = tbl_deg.rows[row_idx]
            replace_cell_text(row.cells[1], fmt0(d["areas"][si]))
            replace_cell_text(row.cells[2], fmt(d["results"][si], 2))
            if si == 0:
                replace_cell_text(row.cells[3], fmt(d["av_r"], 2))
            row_idx += 1

    # ── TBL[136]: LOD ────────────────────────────────────────────────────
    tbl_lod = tbls[32]
    lin_concs = [(l["pct"], l["conc"], l["av"]) for l in lin]
    for i, (pct, conc, av) in enumerate(lin_concs):
        row = tbl_lod.rows[i+2]
        replace_cell_text(row.cells[0], str(i+1))
        replace_cell_text(row.cells[1], str(pct))
        replace_cell_text(row.cells[2], str(conc))
        replace_cell_text(row.cells[3], fmt0(av))
        if i == 0:
            replace_cell_text(row.cells[4], fmt(lod["sd"], 2))
            replace_cell_text(row.cells[5], fmt0(lod["slope"]))
            replace_cell_text(row.cells[6], fmt(lod["ug"], 2))
            replace_cell_text(row.cells[7], fmt0(lod["area"]))

    # ── TBL[153]: LOQ ────────────────────────────────────────────────────
    tbl_loq = tbls[35]
    for i, (pct, conc, av) in enumerate(lin_concs):
        row = tbl_loq.rows[i+2]
        replace_cell_text(row.cells[0], str(i+1))
        replace_cell_text(row.cells[1], str(pct))
        replace_cell_text(row.cells[2], str(conc))
        replace_cell_text(row.cells[3], fmt0(av))
        if i == 0:
            replace_cell_text(row.cells[4], fmt(lod["sd"], 2))
            replace_cell_text(row.cells[5], fmt0(lod["slope"]))
            replace_cell_text(row.cells[6], fmt(loq["ug"], 2))
            replace_cell_text(row.cells[7], fmt0(loq["area"]))

    # ── TBL[039]: Specificity result ──────────────────────────────────────
    tbl_spec = tbls[39]
    replace_bold_in_cell(tbl_spec.rows[1].cells[1], "Avanafil", AI)

    # ── TBL[040]: Observation ─────────────────────────────────────────────
    tbl_obs = tbls[40]
    replace_bold_in_cell(tbl_obs.rows[1].cells[0], "Avanafil", NAME)

    # ── TBL[043]: Conclusion + Statement ──────────────────────────────────
    tbl_con = tbls[43]
    con_cell = tbl_con.rows[1].cells[0]
    for para in con_cell.paragraphs:
        full = "".join(r.text for r in para.runs)
        if "Avanafil Starting Material" in full:
            if is_product:
                replace_bold_in_cell(
                    con_cell, "Avanafil Starting Material", PN)
            else:
                replace_bold_in_cell(
                    con_cell, "Avanafil Starting Material", f"{AI} Starting Material")

   # Statement of Suitability — both occurrences in one pass
    stmt_cell = tbl_con.rows[3].cells[0]
    stmt_name = PN if is_product else AI
    for para in stmt_cell.paragraphs:
        full = "".join(r.text for r in para.runs)
        if "determination of Avanafil" in full and "quantification of Avanafil" in full:
            from lxml import etree
            W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            p_el = para._p
            for r in para.runs:
                p_el.remove(r._r)
            part1, rest = full.split("determination of Avanafil", 1)
            part2, part3 = rest.split("quantification of Avanafil", 1)

            def make_run(text, bold=False):
                r_el = etree.SubElement(p_el, f'{{{W}}}r')
                rpr = etree.SubElement(r_el, f'{{{W}}}rPr')
                sz = etree.SubElement(rpr, f'{{{W}}}sz')
                sz.set(f'{{{W}}}val', '22')
                szcs = etree.SubElement(rpr, f'{{{W}}}szCs')
                szcs.set(f'{{{W}}}val', '22')
                if bold:
                    etree.SubElement(rpr, f'{{{W}}}b')
                    etree.SubElement(rpr, f'{{{W}}}bCs')
                t_el = etree.SubElement(r_el, f'{{{W}}}t')
                t_el.text = text
                t_el.set(
                    '{http://www.w3.org/XML/1998/namespace}space', 'preserve')

            if part1:
                make_run(part1)
            make_run("determination of ", False)
            make_run(AI, True)
            if part2:
                make_run(part2)
            make_run("quantification of ", False)
            make_run(stmt_name, True)
            if part3:
                make_run(part3)

    # ── Generate and insert linearity chart ───────────────────────────────
    try:
        _insert_chart(doc, data)
    except Exception as e:
        print(f"Warning: chart insertion failed: {e}")

    doc.save(output)
    print(f"Saved: {output}")


def _insert_chart(doc, data):
    """Generate linearity chart using matplotlib and replace existing chart in DOCX."""
    from lxml import etree

    AI = data["AI"]
    lin = data["lin"]

    x = [l["pct"] for l in lin]
    y = [float(l["av"]) for l in lin]

    width_in = 4924425 / 914400
    height_in = 2762250 / 914400

    plt.rcParams.update({'font.family': 'Calibri'})
    fig, ax = plt.subplots(figsize=(width_in, height_in), dpi=150)
    fig.patch.set_facecolor('white')
    ax.set_facecolor('white')

    coeffs = np.polyfit(x, y, 1)
    trendline = np.poly1d(coeffs)
    x_line = np.linspace(min(x), max(x), 500)
    ax.plot(x_line, trendline(x_line), color='#4472C4', linewidth=1.2)
    ax.scatter(x, y, color='#4472C4', s=8, zorder=3, linewidths=0)

    ax.set_title(AI, fontsize=11, fontweight='normal', pad=8)
    ax.set_xlabel("Concentration (%)", fontsize=10)
    ax.set_ylabel("Peak Area", fontsize=10)

    ax.set_xlim(60, 140)
    ax.set_xticks(range(60, 141, 10))

    y_max = max(y)
    y_top = int(np.ceil((y_max + 500000) / 1000000)) * 1000000
    ax.set_ylim(4000000, y_top)
    ax.set_yticks(range(4000000, y_top + 1, 1000000))
    ax.yaxis.set_major_formatter(
        plt.FuncFormatter(lambda val, _: f'{int(val)}'))

    ax.tick_params(axis='both', labelsize=8, width=0.4, length=3)
    ax.grid(True, which='major', color='#C0C0C0', linewidth=0.4, linestyle='-')
    ax.set_axisbelow(True)
    for spine in ax.spines.values():
        spine.set_linewidth(0.4)
        spine.set_color('black')
        spine.set_visible(True)

    fig.tight_layout(pad=0.8)

    tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
    tmp_path = tmp.name
    tmp.close()
    fig.savefig(tmp_path, dpi=150, bbox_inches='tight')
    plt.close(fig)

    body = doc.element.body
    C = 'http://schemas.openxmlformats.org/drawingml/2006/chart'
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    chart_drawing = None
    for drawing in body.iter(f'{{{W}}}drawing'):
        for el in drawing.iter():
            if el.tag == f'{{{C}}}chart':
                chart_drawing = drawing
                break
        if chart_drawing is not None:
            break

    if chart_drawing is None:
        print("Warning: chart drawing not found in document")
        os.unlink(tmp_path)
        return

    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.parts.image import ImagePart
    from docx.image.image import Image as DocxImage
    from docx.opc.packuri import PackURI

    with open(tmp_path, 'rb') as f:
        img_bytes = f.read()

    image = DocxImage.from_file(tmp_path)
    image_part = ImagePart(
        PackURI('/word/media/chart_generated.png'),
        'image/png',
        img_bytes,
        image
    )
    rId = doc.part.relate_to(image_part, RT.IMAGE)

    cx = 4924425
    cy = 2762250

    img_xml = f'''<w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <wp:inline xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
             distT="0" distB="0" distL="0" distR="0">
    <wp:extent cx="{cx}" cy="{cy}"/>
    <wp:effectExtent l="0" t="0" r="0" b="0"/>
    <wp:docPr id="1" name="Chart"/>
    <wp:cNvGraphicFramePr/>
    <a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
      <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
        <pic:pic xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">
          <pic:nvPicPr>
            <pic:cNvPr id="1" name="Chart"/>
            <pic:cNvPicPr/>
          </pic:nvPicPr>
          <pic:blipFill>
            <a:blip r:embed="{rId}" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"/>
            <a:stretch><a:fillRect/></a:stretch>
          </pic:blipFill>
          <pic:spPr>
            <a:xfrm><a:off x="0" y="0"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm>
            <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
          </pic:spPr>
        </pic:pic>
      </a:graphicData>
    </a:graphic>
  </wp:inline>
</w:drawing>'''

    new_drawing = etree.fromstring(img_xml)
    parent = chart_drawing.getparent()
    idx = list(parent).index(chart_drawing)
    parent.remove(chart_drawing)
    parent.insert(idx, new_drawing)

    os.unlink(tmp_path)
    print(f"Chart inserted: {AI}, {len(x)} points")


def _update_chart(doc, areas):
    """Legacy — no longer called."""
    pass


# ── Main ───────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    BASE = os.path.dirname(os.path.abspath(__file__))
    excel = sys.argv[1] if len(sys.argv) > 1 else os.path.join(
        BASE, "AMV_report.xlsx")
    output = sys.argv[2] if len(sys.argv) > 2 else os.path.join(
        BASE, "Method_Validation_Report.docx")
    template = os.path.join(BASE, "approved_file.docx")

    print(f"Reading: {excel}")
    data = read_excel(excel)
    print(f"Product: {data['AI']} | Report: {data['RPT_NO']}")
    fill_template(data, template, output)
